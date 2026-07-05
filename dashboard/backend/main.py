import os
import pymysql
import pandas as pd
import logging
import sys
from fastapi import FastAPI, Query, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from typing import Optional, List
from collector.builder import build_from_api, build_from_csv
import google.generativeai as genai
from docx import Document
from io import BytesIO
from fastapi.responses import StreamingResponse
import json
import datetime
import subprocess

# 로깅 설정: 시간 정보를 포함하도록 포맷 지정
LOG_FORMAT = '%(asctime)s %(levelname)s: %(message)s'
DATE_FORMAT = '%Y-%m-%d %H:%M:%S'

logging.basicConfig(
    level=logging.INFO,
    format=LOG_FORMAT,
    datefmt=DATE_FORMAT,
    stream=sys.stdout
)

# uvicorn의 기본 로깅 설정을 덮어씌워 시간 정보가 표시되도록 함
for name in ["uvicorn", "uvicorn.error", "uvicorn.access"]:
    log = logging.getLogger(name)
    log.handlers = []  # 기존 핸들러 제거
    handler = logging.StreamHandler(sys.stdout)
    handler.setFormatter(logging.Formatter(LOG_FORMAT, datefmt=DATE_FORMAT))
    log.addHandler(handler)
    log.propagate = False

logger = logging.getLogger(__name__)

# Gemini Config
api_key = os.environ.get("GEMINI_API_KEY")
if api_key:
    genai.configure(api_key=api_key)
    logger.info("Gemini API configured.")
else:
    logger.warning("GEMINI_API_KEY not found. Report generation will fail.")

app = FastAPI()

def get_conn():
    return pymysql.connect(
        host="mariadb",
        user="root",
        password="password",
        database="ai_lawsuits",
        cursorclass=pymysql.cursors.DictCursor
    )

def get_data_dir():
    """Helper to find the data directory robustly using absolute paths"""
    # This file is in backend/main.py, so project root is one level up
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    data_dir = os.path.join(base_dir, "data")
    return data_dir

def find_col_in_df(keys, df):
    """Helper to find a column in a dataframe based on multiple potential names"""
    # 1. Try exact match first
    for k in keys:
        for col in df.columns:
            if k.lower() == col.lower():
                return col
    
    # 2. Try partial match, but avoid columns containing both '원고' and '피고' if we are looking for only one
    for k in keys:
        for col in df.columns:
            if k.lower() in col.lower():
                # If searching for '피고', avoid '원고 v. 피고'
                if k == "피고" and "원고" in col:
                    continue
                if k == "원고" and "피고" in col:
                    continue
                return col
    return None

@app.get("/api/cases")
def get_cases(file_name: Optional[str] = Query(None, description="Name of the CSV file in ./data/ to use")):
    """
    Main endpoint to fetch cases. 
    1. If file_name is provided and exists in ./data/, load from CSV.
    2. Otherwise, fetch live data from CourtListener API.
    """
    if file_name:
        csv_path = os.path.join(get_data_dir(), file_name)
        if os.path.exists(csv_path):
            try:
                data = build_from_csv(csv_path)
                return {"source": "csv", "file": file_name, "count": len(data), "data": data}
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Error processing CSV: {str(e)}")
        else:
            raise HTTPException(status_code=404, detail=f"File {file_name} not found in ./data/")
    
    # Default: Fetch from API
    try:
        data = build_from_api()
        return {"source": "api", "query": "artificial intelligence", "count": len(data), "data": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching from API: {str(e)}")

@app.get("/api/version")
def get_version():
    """
    Returns the current application version based on git tags and commit hash.
    """
    try:
        import shutil
        if not shutil.which('git'):
            return {"version": "git-not-found", "commit": "unknown"}

        # Avoid 'dubious ownership' error in Docker
        subprocess.run(["git", "config", "--global", "--add", "safe.directory", "/app"], check=False)
        
        # Get the latest tag
        try:
            version = subprocess.check_output(["git", "tag", "--sort=-v:refname"], 
                                             stderr=subprocess.STDOUT, 
                                             cwd="/app").decode("utf-8").split('\n')[0].strip()
            if not version:
                 version = subprocess.check_output(["git", "describe", "--tags", "--always"], 
                                                  stderr=subprocess.STDOUT, 
                                                  cwd="/app").decode("utf-8").strip()
        except:
            version = subprocess.check_output(["git", "describe", "--tags", "--always"], 
                                              stderr=subprocess.STDOUT, 
                                              cwd="/app").decode("utf-8").strip()
            
        commit = subprocess.check_output(["git", "rev-parse", "HEAD"], 
                                         stderr=subprocess.STDOUT, 
                                         cwd="/app").decode("utf-8").strip()
        
        # Get commit date in "May 3 19:14:33 2026" format
        try:
            # Use LC_ALL=C to ensure English month names
            env = os.environ.copy()
            env["LC_ALL"] = "C"
            date = subprocess.check_output(["git", "show", "-s", "--format=%cd", "--date=format:%b %-d %H:%M:%S %Y", "HEAD"], 
                                          stderr=subprocess.STDOUT, 
                                          cwd="/app",
                                          env=env).decode("utf-8").strip()
        except:
            date = "unknown-date"
            
        return {"version": version, "commit": commit, "date": date}
    except Exception as e:
        logger.error(f"Error getting git version: {str(e)}")
        # Return error info in version field to help debugging
        return {"version": f"error: {str(e)[:50]}", "commit": "unknown", "date": "unknown"}

@app.get("/api/db-cases")
def get_db_cases():
    """Fallback to original database functionality with proper connection closing"""
    conn = get_conn()
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM lawsuits ORDER BY file_date DESC")
        data = cur.fetchall()
        return {"source": "database", "count": len(data), "data": data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")
    finally:
        conn.close()

@app.get("/api/files")
def list_data_files():
    """Helper to see available CSVs"""
    data_dir = get_data_dir()
    if not os.path.exists(data_dir):
        return []
    # Return files sorted by name descending (latest first)
    return sorted([f for f in os.listdir(data_dir) if f.endswith(".csv")], reverse=True)

@app.get("/api/statistics")
def get_statistics(file_name: Optional[str] = Query(None)):
    """
    Returns aggregated statistics for the dashboard.
    """
    if file_name:
        csv_path = os.path.join(get_data_dir(), file_name)
        if os.path.exists(csv_path):
            # Skip first 2 lines based on format: 1. Title, 2. Extraction Date
            df = pd.read_csv(csv_path, skiprows=2).fillna("")
            # Robust filtering: Ensure only rows with a valid case title are kept
            title_col = find_col_in_df(["소송제목 (원고 v. 피고)*", "소송제목 (원고 v. 피고)", "Case Name"], df)
            if title_col:
                df = df[df[title_col].astype(str).str.strip() != ""]
        else:
            raise HTTPException(status_code=404, detail="File not found")
    else:
        # Fallback to default
        data = build_from_api()
        df = pd.DataFrame(data)

    def get_top_n(series, n=15, key_type=None):
        # Convert to string and strip whitespace to ensure clean grouping
        series = series.astype(str).str.strip()
        # Filter out empty strings or "Unknown"
        series = series[series.str.lower() != "unknown"]
        series = series[series != ""]
        
        # Capture the actual number of cases (rows) after filtering unknowns
        # but before any category-specific expansion (like exploding defendants)
        case_count = len(series)

        if key_type == "defendant":
            # Support multiple defendants per case by splitting by comma
            series = series.str.split(',').explode().str.strip()
            # Basic cleaning for defendant names
            def clean_def(name):
                # Remove "et al." or "et al"
                name = name.replace("et al.", "").replace("et al", "").strip()
                return name
            series = series.apply(clean_def)
            # Re-filter after splitting in case empty segments were created
            series = series[series != ""]

        if key_type == "target_data":
            # Categorize Target Data into requested groups
            def categorize_data(text):
                text = text.lower()
                mapping = {
                    "Authors/Books": ["book", "text", "novel", "author", "책", "도서", "저서", "글", "문학", "작가"],
                    "Music/Recordings": ["music", "audio", "song", "recording", "음악", "음원", "노래", "오디오", "가수"],
                    "Visual Art": ["image", "art", "picture", "photo", "drawing", "painting", "이미지", "그림", "사진", "미술", "예술", "작품"],
                    "Video/Film": ["video", "movie", "film", "youtube", "비디오", "영상", "영화", "유튜브"],
                    "Voice/Likeness": ["voice", "face", "likeness", "biometric", "목소리", "얼굴", "초상", "음성"],
                    "News": ["news", "article", "journalism", "뉴스", "기사", "언론"],
                    "Publishers": ["publish", "publisher", "press", "출판", "신문사"],
                    "Code/Software": ["code", "software", "programming", "github", "코드", "소프트웨어", "프로그래밍", "깃허브"]
                }
                for cat, keywords in mapping.items():
                    if any(kw in text for kw in keywords):
                        return cat
                return "Other"
            series = series.apply(categorize_data)

        if key_type == "claim":
            # Categorize Claims into requested groups
            def categorize_claim(text):
                text = text.lower()
                mapping = {
                    "Direct copyright": ["direct", "copyright infringement", "저작권 침해", "복제", "무단 사용", "무단 수집"],
                    "Vicarious infringement": ["vicarious", "대위"],
                    "Contributory infringement": ["contributory", "기여", "방조"],
                    "Inducement": ["inducement", "유도", "권장"],
                    "DMCA": ["dmca", "digital millennium", "저작권 보호"],
                    "Latham Act": ["latham", "lanham", "상표", "부정경쟁"],
                    "Sate law claims": ["state law", "common law", "주법", "관습법", "불공정", "퍼블리시티"]
                }
                for cat, keywords in mapping.items():
                    if any(kw in text for kw in keywords):
                        return cat
                return "기타"
            series = series.apply(categorize_claim)

        return series.value_counts().head(n).to_dict(), case_count

    col_mapping = {
        "claim": ["청구 내용", "Reason", "소송 사유", "Claim", "소송이유"],
        "defendant": ["피고", "Defendant", "피고*"],
        "country": ["국가", "Country", "국가*"],
        "decision": ["진행현황", "Status", "진행현황*", "Decision"],
        "target_data": ["대상 데이터", "Target Data", "대상 데이터*", "대상 데이타"]
    }

    stats = {}
    for key, keys in col_mapping.items():
        col = find_col_in_df(keys, df)
        if col:
            data_dict, total_count = get_top_n(df[col], key_type=key)
            stats[key] = {"data": data_dict, "total": total_count}
        else:
            stats[key] = {"data": {}, "total": 0}

    return stats

@app.post("/api/report/generate")
def generate_report(request: dict):
    """
    Generates a monthly report using Gemini.

    NOTE: 동기(`def`) 엔드포인트로 선언해 FastAPI가 스레드풀에서 실행하도록 한다.
    내부의 `model.generate_content()`는 수십 초가 걸릴 수 있는 블로킹 호출이라,
    `async def`로 두면 이벤트 루프가 통째로 멈춰 다른 요청 처리가 막히고
    연결이 끊겨 클라이언트에서 NetworkError가 발생할 수 있다.
    """
    report_type = request.get("type")
    month = request.get("month") # YYYY-MM
    file_name = request.get("file_name")

    if not api_key:
        raise HTTPException(status_code=500, detail="Gemini API Key not configured")

    if file_name:
        csv_path = os.path.join(get_data_dir(), file_name)
        if os.path.exists(csv_path):
            # Skip first 2 lines based on format: 1. Title, 2. Extraction Date
            df = pd.read_csv(csv_path, skiprows=2).fillna("")
            # Filter empty rows
            title_col_filter = find_col_in_df(["소송제목 (원고 v. 피고)*", "소송제목 (원고 v. 피고)", "Case Name"], df)
            if title_col_filter:
                df = df[df[title_col_filter].astype(str).str.strip() != ""]
        else:
            raise HTTPException(status_code=404, detail="File not found")
    else:
        data = build_from_api()
        df = pd.DataFrame(data)

    # Filter by month
    if report_type == "filing_date":
        date_col = find_col_in_df(["소송제기일", "Filing Date"], df)
    else:
        # Use explicit column name provided in the mission
        date_col = find_col_in_df(["Last Update(업로드 시 제외)", "Last Update", "최근 업데이트"], df)

    if not date_col:
        raise HTTPException(status_code=400, detail=f"Date column for {report_type} not found")

    df[date_col] = pd.to_datetime(df[date_col], errors='coerce')
    filtered_df = df[df[date_col].dt.strftime('%Y-%m') == month]

    if filtered_df.empty:
        return {"report": f"### {month} 월간 소송 현황 보고서\n\n해당 월에 해당하는 데이터가 없습니다."}

    # Prepare summary data for the full database (for cumulative stats)
    total_cases_all = len(df)
    defendant_col = find_col_in_df(["피고", "Defendant", "피고*"], df)
    top_5_defendants = df[defendant_col].value_counts().head(5).to_dict() if defendant_col else {}
    
    yearly_trend = df[date_col].dt.year.value_counts().sort_index().to_dict()

    # Prepare detailed data for the target month
    case_summary = filtered_df.to_json(orient="records", force_ascii=False)
    
    prompt = f"""
첨부한 데이터는 {month} 기준 AI 데이터 소송 관련 데이터베이스입니다. 
이 데이터를 바탕으로 전문적인 'AI 데이터 소송 월간 동향 보고서'를 작성해 주세요. 

[제공된 데이터 - {month} 상세]
{case_summary}

[전체 데이터베이스 통계 정보]
- 전체 누적 소송 건수: {total_cases_all}
- 누적 피고 순위 Top 5: {top_5_defendants}
- 연도별 소송 증가 추이: {yearly_trend}

보고서에는 다음 항목이 반드시 포함되어야 합니다:

1. Executive Summary:
   - 해당 월({month})의 신규 소송 건수 요약.
   - 가장 핵심적인 법적 트렌드(예: TPM 우회, 저작권 침해, 무단 스크래핑 등) 요약.

2. 신규 소송 현황 분석:
   - 신규 소송의 총 건수, 주요 피고 기업, 국가별 분포를 표(Table)로 정리.
   - 소송 목록을 날짜순으로 정리하고, 각 소송의 핵심 쟁점(핵심 논거)을 기술.

3. 유형별 심층 분석:
   - 데이터에서 가장 높은 비중을 차지하는 소송 유형(예: YouTube 무단 스크래핑, 개인정보 침해 등)을 그룹화하여 분석.
   - 특정 기업(예: OpenAI, Google, Runway AI 등)에 소송이 집중된 경우 그 원인과 증거(내부 문서 유출 등)를 상세히 기술.

4. 글로벌 동향 및 특이 사례:
   - 미국 외 국가(유럽 GDPR, 한국 방송사 소송 등)에서 발생한 주요 소송의 의미 분석.
   - 퍼블리시티권이나 AI 음성 합성 같은 신규 소송 유형이 있다면 별도로 강조.

5. 통계 및 전망:
   - 전체 데이터베이스를 기준으로 한 연도별 소송 증가 추이와 누적 피고 순위 Top 5 도출.
   - 향후 법적 판결이 업계에 미칠 영향이나 주목해야 할 모니터링 포인트(Monitoring Points) 제시.

[보고서 형식 및 지침]
- 가독성을 위해 표(Table), 불렛 포인트, 섹션 구분을 명확히 사용하세요.
- 톤앤매너는 전문적인 기술 전략 보고서 스타일로 작성하세요.
- 한국어로 작성하고 마크다운(Markdown) 형식을 유지하세요.
- 별도의 대제목(## {month} 보고서 등)은 생략하고 바로 '1. Executive Summary'부터 시작하세요.
"""

    try:
        # Safety settings matching reference
        safety_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
        ]

        model = genai.GenerativeModel(
            model_name='gemini-flash-latest',
            safety_settings=safety_settings
        )
        response = model.generate_content(prompt)
        
        if response.candidates and response.candidates[0].content.parts:
            logger.info(f"Report generated successfully for {month} ({report_type})")
            return {"report": response.text}
        else:
            logger.error("Gemini response blocked or no candidates. This might be due to content safety filters.")
            return {"report": "### 보고서 생성 실패\n\nGemini AI의 응답이 차단되었습니다. 데이터의 민감도 설정을 확인하거나, 질문 내용을 검토해주세요."}
            
    except Exception as e:
        logger.error(f"Gemini API Error: {str(e)}")
        if not api_key:
            logger.critical("CRITICAL: GEMINI_API_KEY is missing. Monthly report generation is disabled.")
            raise HTTPException(status_code=500, detail="GEMINI_API_KEY가 설정되지 않았습니다. 깃허브 레포 설정 또는 서버 환경변수를 확인하세요.")
        raise HTTPException(status_code=500, detail=f"Gemini API Error: {str(e)}")

@app.post("/api/report/download")
async def download_report(request: dict):
    """
    Converts markdown report to docx and returns for download.
    """
    report_content = request.get("content")
    title = request.get("title", "Monthly_Litigation_Report")
    
    doc = Document()
    doc.add_heading(title, 0)
    
    lines = report_content.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line.lstrip('#').strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line.lstrip('#').strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.lstrip('#').strip(), level=3)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line)

    target_stream = BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    
    filename = f"{title}_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
    
    return StreamingResponse(
        target_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# Mount frontend static files
# 프로젝트 루트(= dashboard/)를 main.py 위치 기준 절대경로로 계산한다.
# 상대경로(directory="frontend/css")는 실행 위치(CWD)에 의존해, 모노레포로
# 폴더를 옮긴 뒤 다른 위치에서 실행하면 정적 파일이 404가 날 수 있다.
# → 모든 정적 마운트/응답을 _base_dir 기준 절대경로로 통일해 CWD와 무관하게 동작.
_base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_frontend_dir = os.path.join(_base_dir, "frontend")

app.mount("/css", StaticFiles(directory=os.path.join(_frontend_dir, "css")), name="css")
app.mount("/js", StaticFiles(directory=os.path.join(_frontend_dir, "js")), name="js")
app.mount("/assets", StaticFiles(directory=os.path.join(_frontend_dir, "assets")), name="assets")
app.mount("/img", StaticFiles(directory=os.path.join(_base_dir, "img")), name="img")
app.mount("/timeline", StaticFiles(directory=os.path.join(_base_dir, "timeline")), name="timeline")
# 매뉴얼 문서(api_guide.md / user_guide.md 등)는 frontend/manual/ 에 있다.
# 이 마운트가 없으면 /manual/* 요청이 404 → "Documentation file not found"로 실패한다.
app.mount("/manual", StaticFiles(directory=os.path.join(_frontend_dir, "manual")), name="manual")

@app.get("/")
def read_root():
    # 첫 화면 = 메인 현황판(전체 현황) → 세부 분석은 지도 대시보드(/map.html)로 이동
    return FileResponse(os.path.join(_frontend_dir, "overview.html"))

@app.get("/map.html")
def read_map():
    return FileResponse(os.path.join(_frontend_dir, "index.html"))

@app.get("/lineage.html")
def read_lineage():
    return FileResponse(os.path.join(_frontend_dir, "lineage.html"))

@app.get("/overview.html")
def read_overview():
    return FileResponse(os.path.join(_frontend_dir, "overview.html"))
